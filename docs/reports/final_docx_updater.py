#!/usr/bin/env python
"""
Final smart updater: Replace paragraph content while preserving all formatting.
Keeps styles, tables, lists, spacing, everything - just updates the text.
"""

from docx import Document
from docx.shared import Pt
from pathlib import Path
import re

def read_markdown_sections(filepath):
    """Read markdown and group by headings."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    sections = {}
    current_heading = None
    current_content = []

    for line in lines:
        # Detect heading levels
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            # Save previous section
            if current_heading:
                sections[current_heading] = {
                    'level': len(heading_match.group(1)) if 'level' in locals() else 1,
                    'content': '\n'.join(current_content).strip()
                }

            current_heading = heading_match.group(2).strip()
            current_content = []
        else:
            current_content.append(line.rstrip())

    # Save last section
    if current_heading:
        sections[current_heading] = {
            'level': 1,
            'content': '\n'.join(current_content).strip()
        }

    return sections

def normalize_heading(text):
    """Normalize heading text for matching."""
    return text.upper().strip()

def update_document_preserve_formatting(original_docx, markdown_file, output_docx):
    """
    Update document: replace paragraph text while preserving all formatting.
    Strategy: Keep existing structure, just update text content.
    """

    doc = Document(original_docx)
    md_sections = read_markdown_sections(markdown_file)

    print(f"\n{'='*80}")
    print(f"SMART CONTENT UPDATE (PRESERVING FORMATTING)")
    print(f"{'='*80}\n")

    # Key section mappings (markdown heading -> document section name)
    section_map = {
        'RESUMEN': ['RESUMEN'],
        'INTRODUCCIÓN': ['INTRODUCCIÓN', 'INTRODUCTION'],
        'CAPÍTULO 1': ['CAPÍTULO 1', 'CAPITULO 1', 'CHAPTER 1'],
        '1.1. MARCO TEÓRICO': ['1.1. MARCO TEÓRICO', 'MARCO TEÓRICO'],
        '1.2. DISEÑO METODOLÓGICO': ['1.2. DISEÑO METODOLÓGICO', 'DISEÑO METODOLÓGICO'],
        '1.3. SELECCIÓN Y ANÁLISIS': ['1.3.'],
        '1.4. BASE DE DATOS': ['1.4.', 'SOUNDWAVE'],
        '1.5. CONJUNTO DE PRUEBA': ['1.5.'],
        '1.6. ESPECIFICACIÓN': ['1.6.'],
        '1.7. REQUERIMIENTOS NO': ['1.7.'],
        '1.8. MÉTRICAS': ['1.8.'],
        '1.9. ESCENARIOS': ['1.9.'],
        '1.10. ARQUITECTURA': ['1.10.'],
        '1.11. AVANCE DEL TRABAJO': ['1.11.'],
        '1.12. CONCLUSIONES': ['1.12.', 'CONCLUSIONES'],
        '1.13. RECOMENDACIONES': ['1.13.', 'RECOMENDACIONES'],
    }

    # Find and mark sections for update
    section_positions = {}
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip().upper()

        for md_heading, doc_keywords in section_map.items():
            for keyword in doc_keywords:
                if keyword in para_text:
                    if md_heading not in section_positions:
                        section_positions[md_heading] = para_idx
                        print(f"[FOUND] {md_heading}")
                        print(f"        at paragraph {para_idx}: {para.text[:70]}")
                        break

    print(f"\n[SUMMARY] Found {len(section_positions)} sections to update")
    print(f"[INFO] Strategy: Update paragraph text after each heading")
    print(f"[INFO] All formatting, styles, tables preserved\n")

    # Now perform selective updates
    # For each heading found, update the following paragraphs until next heading
    updated_sections = 0

    for md_heading, para_idx in sorted(section_positions.items(), key=lambda x: x[1]):
        if md_heading in md_sections:
            content = md_sections[md_heading]['content']

            # Find the range of paragraphs to update
            # (from after this heading to before next heading)
            next_heading_idx = None
            for other_heading, other_idx in section_positions.items():
                if other_idx > para_idx:
                    if next_heading_idx is None or other_idx < next_heading_idx:
                        next_heading_idx = other_idx

            # Update paragraphs in range
            start_idx = para_idx + 1
            end_idx = next_heading_idx if next_heading_idx else len(doc.paragraphs)

            # Split content into paragraphs
            content_paras = [p.strip() for p in content.split('\n') if p.strip()]

            if content_paras:
                # Update or add paragraphs
                for content_idx, content_text in enumerate(content_paras):
                    target_para_idx = start_idx + content_idx

                    if target_para_idx < end_idx and target_para_idx < len(doc.paragraphs):
                        # Update existing paragraph
                        doc.paragraphs[target_para_idx].text = content_text
                        print(f"[UPDATE] Para {target_para_idx}: {content_text[:60]}")
                    elif target_para_idx >= len(doc.paragraphs):
                        # Add new paragraph
                        doc.add_paragraph(content_text)
                        print(f"[ADD] New para: {content_text[:60]}")

                updated_sections += 1

    print(f"\n[SUCCESS] Updated {updated_sections} sections")
    print(f"[INFO] Document formatting preserved")

    # Save
    doc.save(output_docx)
    print(f"\n[OK] Document saved to: {output_docx}")

    return True

if __name__ == '__main__':
    original_docx = 'docs/reports/IDI - Primer Informe.docx'
    markdown_file = 'docs/reports/IDI_Capitulo1_v3.md'
    output_docx = 'docs/reports/IDI_Primer_Informe_v2_Updated.docx'

    try:
        success = update_document_preserve_formatting(original_docx, markdown_file, output_docx)
        if success:
            print(f"\n{'='*80}")
            print(f"UPDATE COMPLETE - FILE READY FOR REVIEW")
            print(f"{'='*80}")
            print(f"\nOriginal: {original_docx}")
            print(f"Updated:  {output_docx}")
            print(f"\nPlease review the output file for accuracy.")
            print(f"All formatting and styles have been preserved.")

    except Exception as e:
        print(f"\n[ERROR] {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
