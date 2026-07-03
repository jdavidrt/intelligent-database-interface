#!/usr/bin/env python
"""
Analyze Word document structure and intelligently update with markdown content.
Preserves all formatting, styles, tables, and document structure.
"""

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from pathlib import Path
import re

def analyze_document_structure(docx_filepath):
    """Analyze and print document structure."""
    doc = Document(docx_filepath)

    print(f"\n{'='*80}")
    print(f"DOCUMENT STRUCTURE ANALYSIS: {Path(docx_filepath).name}")
    print(f"{'='*80}")

    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")

    print(f"\n--- FIRST 30 PARAGRAPHS (to understand structure) ---\n")

    for idx, para in enumerate(doc.paragraphs[:30]):
        if para.text.strip():
            style = para.style.name if para.style else "No Style"
            level = para.paragraph_format.outline_level if hasattr(para.paragraph_format, 'outline_level') else "N/A"
            text_preview = para.text[:70]
            print(f"{idx:3d} | Style: {style:20s} | Level: {str(level):2s} | {text_preview}")

    return len(doc.paragraphs), len(doc.tables)

def read_markdown_with_structure(filepath):
    """Read markdown preserving its structure."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    sections = []
    current_section = {
        'heading': '',
        'level': 0,
        'content': []
    }

    for line in lines:
        # Count heading level
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            # Save previous section if it has content
            if current_section['heading'] or current_section['content']:
                sections.append(current_section)

            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            current_section = {
                'heading': heading_text,
                'level': level,
                'content': []
            }
        else:
            current_section['content'].append(line)

    # Add last section
    if current_section['heading'] or current_section['content']:
        sections.append(current_section)

    return sections

def update_document_smart(original_docx, markdown_file, output_docx):
    """
    Smart update: analyze document, parse markdown, then carefully merge.
    """
    doc = Document(original_docx)
    md_sections = read_markdown_with_structure(markdown_file)

    print(f"\n{'='*80}")
    print(f"MARKDOWN STRUCTURE")
    print(f"{'='*80}")
    print(f"\nFound {len(md_sections)} sections in markdown:\n")

    for idx, section in enumerate(md_sections[:15]):
        if section['heading']:
            indent = "  " * (section['level'] - 1)
            content_lines = len([l for l in section['content'] if l.strip()])
            print(f"{indent}[H{section['level']}] {section['heading']} ({content_lines} content lines)")

    # Strategy: Find sections by heading matching and update paragraph content after them
    heading_keywords = {
        'RESUMEN': 'summary',
        'INTRODUCCIÓN': 'intro',
        'CAPÍTULO 1': 'chapter1',
        'MARCO TEÓRICO': 'framework',
        'CONCLUSIONES': 'conclusions',
        'RECOMENDACIONES': 'recommendations',
    }

    print(f"\n{'='*80}")
    print(f"MATCHING STRATEGY")
    print(f"{'='*80}\n")

    matched_headings = []
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_upper = paragraph.text.strip().upper()

        for keyword in heading_keywords.keys():
            if keyword in para_upper:
                matched_headings.append({
                    'keyword': keyword,
                    'para_idx': para_idx,
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else "Unknown"
                })
                print(f"MATCH at paragraph {para_idx}: '{paragraph.text[:60]}...'")
                print(f"       Style: {paragraph.style.name if paragraph.style else 'Unknown'}")

    print(f"\n[SUMMARY] Found {len(matched_headings)} matching section headings in document")
    print(f"\n[OK] Document structure is intact and ready for targeted updates")
    print(f"[OK] Next step: Implement paragraph content replacement")

    # Save analysis (preserve original, create new file)
    doc.save(output_docx)
    print(f"\n[SUCCESS] Document saved (unchanged) to: {output_docx}")
    print(f"[INFO] All formatting and styles preserved")

    return matched_headings

if __name__ == '__main__':
    original_docx = 'docs/reports/IDI - Primer Informe.docx'
    markdown_file = 'docs/reports/IDI_Capitulo1_v3.md'
    output_docx = 'docs/reports/IDI_Primer_Informe_v2_Updated.docx'

    try:
        # Step 1: Analyze document structure
        para_count, table_count = analyze_document_structure(original_docx)

        # Step 2: Read and analyze markdown
        # Step 3: Match sections and prepare update
        matched = update_document_smart(original_docx, markdown_file, output_docx)

        print(f"\n{'='*80}")
        print(f"ANALYSIS COMPLETE")
        print(f"{'='*80}")
        print(f"\nDocument Details:")
        print(f"  - Paragraphs: {para_count}")
        print(f"  - Tables: {table_count}")
        print(f"  - Matched sections: {len(matched)}")
        print(f"\nNext: Implement targeted content updates for each section")

    except Exception as e:
        print(f"\n[ERROR] {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
