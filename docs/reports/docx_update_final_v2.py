#!/usr/bin/env python
"""
Structure-aware Word document updater.

Syncs the content of IDI_Capitulo1_v3.md into "IDI - Primer Informe.docx",
preserving all existing formatting (paragraph/run styles, table styles,
embedded images). Unlike a naive paragraph-index copy, this script:

  - Locates each section by its heading (numbered "1.6." style or literal
    RESUMEN/INTRODUCCION/REFERENCIAS), independent of position drift.
  - Reuses existing paragraphs' formatting (keeps the first run, clears
    extras) when overwriting text, so styles are not lost.
  - Creates new paragraphs by cloning a template paragraph's XML when the
    markdown has more content than the docx currently holds (e.g. the new
    OE1-OE4 list, the chapter-opening paragraph, section 1.13, References).
  - Deletes surplus paragraphs when the docx has more than the markdown.
  - Updates table cells in place via header-name column matching (so an
    extra/reordered markdown column, e.g. the "#" column in the Brechas
    table, is safely dropped instead of desyncing every other column).
  - Never touches paragraphs that contain embedded images/drawings (the
    real sandbox screenshots living in section 1.11); markdown
    "[ESPACIO PARA CAPTURA DE PANTALLA: ...]" placeholder lines are used
    only as segment separators and are otherwise discarded.
"""

import copy
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

NUM_HEADING_RE = re.compile(r'^(\d+(?:\.\d+)*)\.\s*(.*)$')
PLACEHOLDER_RE = re.compile(r'^\[ESPACIO PARA CAPTURA DE PANTALLA:?\s*(.*)\]$', re.IGNORECASE)
RULE_RE = re.compile(r'^[─\-]{5,}$')
TOP_LEVEL_KEYS = ('RESUMEN', 'INTRODUCCIÓN', 'REFERENCIAS')


# --------------------------------------------------------------------------
# Markdown parsing
# --------------------------------------------------------------------------

class Heading:
    def __init__(self, key, raw_text):
        self.key = key
        self.raw_text = raw_text
        self.blocks = []  # list of ('para', text) | ('tablerow', cells) | ('image', caption)


def parse_markdown(md_path):
    text = Path(md_path).read_text(encoding='utf-8')
    lines = text.split('\n')

    # the file ends with a university footer/signature block after the final
    # rule separator -- that's page-footer content, not part of REFERENCIAS
    last_rule_idx = max(
        (i for i, l in enumerate(lines) if RULE_RE.match(l.strip())),
        default=None,
    )

    headings = []
    current = None
    in_fence = False
    started = False

    for line_idx, raw_line in enumerate(lines):
        stripped = raw_line.strip()

        if last_rule_idx is not None and line_idx >= last_rule_idx:
            break

        if not started:
            if stripped == 'RESUMEN':
                current = Heading('RESUMEN', stripped)
                headings.append(current)
                started = True
            continue

        if not stripped:
            continue
        if RULE_RE.match(stripped):
            continue
        if stripped.startswith('```'):
            in_fence = not in_fence
            continue
        if in_fence:
            current.blocks.append(('para', raw_line.rstrip()))
            continue
        if stripped in TOP_LEVEL_KEYS:
            current = Heading(stripped, stripped)
            headings.append(current)
            continue
        if stripped.upper().startswith('CAPÍTULO '):
            current = Heading('CAPITULO1', stripped)
            headings.append(current)
            continue
        m = NUM_HEADING_RE.match(stripped)
        if m:
            current = Heading(m.group(1), stripped)
            headings.append(current)
            continue
        if stripped.startswith('|'):
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue  # markdown separator row
            current.blocks.append(('tablerow', cells))
            continue
        ph = PLACEHOLDER_RE.match(stripped)
        if ph:
            current.blocks.append(('image', ph.group(1)))
            continue
        current.blocks.append(('para', stripped))

    # group consecutive tablerow entries into single table blocks
    for h in headings:
        grouped = []
        buf = []
        for b in h.blocks:
            if b[0] == 'tablerow':
                buf.append(b[1])
            else:
                if buf:
                    grouped.append(('table', buf))
                    buf = []
                grouped.append(b)
        if buf:
            grouped.append(('table', buf))
        h.blocks = grouped

    return headings


# --------------------------------------------------------------------------
# Text helpers
# --------------------------------------------------------------------------

def normalize(s):
    s = unicodedata.normalize('NFKD', s).encode('ascii', 'ignore').decode('ascii')
    return s.strip().lower()


def paragraph_has_image(paragraph):
    return bool(paragraph._p.findall('.//' + qn('w:drawing'))) or \
        bool(paragraph._p.findall('.//' + qn('w:pict')))


def rebuild_paragraph_text(paragraph, text):
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for extra in runs[1:]:
            extra._r.getparent().remove(extra._r)
    else:
        paragraph.add_run(text)


def clone_paragraph(template_paragraph, text):
    new_el = copy.deepcopy(template_paragraph._p)
    new_p = Paragraph(new_el, template_paragraph._parent)
    rebuild_paragraph_text(new_p, text)
    return new_p


def set_cell_text(cell, text):
    paragraphs = cell.paragraphs
    rebuild_paragraph_text(paragraphs[0], text)
    for extra_p in paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)


def update_table(table, md_rows, log):
    docx_header = [c.text.strip() for c in table.rows[0].cells]
    md_header = md_rows[0]

    col_map = [None] * len(docx_header)
    used_md_cols = set()
    for j, dh in enumerate(docx_header):
        for mi, mh in enumerate(md_header):
            if mi in used_md_cols:
                continue
            if normalize(mh) == normalize(dh):
                col_map[j] = mi
                used_md_cols.add(mi)
                break
    for j, dh in enumerate(docx_header):
        if col_map[j] is not None:
            continue
        for mi, mh in enumerate(md_header):
            if mi in used_md_cols:
                continue
            if normalize(dh) in normalize(mh) or normalize(mh) in normalize(dh):
                col_map[j] = mi
                used_md_cols.add(mi)
                break
    # positional fallback: a docx column with no name match paired with the
    # md column at the same index, if that md column is itself still unused
    # (handles a straight header rename, e.g. "Benchmark" -> "Fuente de patrones")
    for j, dh in enumerate(docx_header):
        if col_map[j] is None and j < len(md_header) and j not in used_md_cols:
            col_map[j] = j
            used_md_cols.add(j)
            log.append(f"    [INFO] table col '{dh}' has no name match; paired positionally with md col {j} ({md_header[j]!r})")

    if any(c is None for c in col_map):
        log.append(f"    [WARN] table col unmatched: docx headers={docx_header} md headers={md_header}")

    # grow/shrink row count to match markdown before writing cell text
    target_rows = len(md_rows)
    current_rows = len(table.rows)
    if target_rows > current_rows:
        last_tr = table.rows[-1]._tr
        for _ in range(target_rows - current_rows):
            new_tr = copy.deepcopy(last_tr)
            last_tr.addnext(new_tr)
            last_tr = new_tr
        log.append(f"    [INFO] table grew from {current_rows} to {target_rows} rows")
    elif target_rows < current_rows:
        for row in table.rows[target_rows:]:
            row._tr.getparent().remove(row._tr)
        log.append(f"    [INFO] table shrank from {current_rows} to {target_rows} rows")

    rows = table.rows
    for r in range(len(rows)):
        docx_cells = rows[r].cells
        md_row = md_rows[r]
        for ci, mi in enumerate(col_map):
            if mi is not None and mi < len(md_row):
                set_cell_text(docx_cells[ci], md_row[mi])


# --------------------------------------------------------------------------
# Body traversal
# --------------------------------------------------------------------------

def build_body_index(doc):
    """Ordered list of {'type': 'p'|'tbl', 'obj': Paragraph|Table} for direct body children."""
    body = doc.element.body
    items = []
    pi = ti = 0
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            items.append({'type': 'p', 'obj': doc.paragraphs[pi]})
            pi += 1
        elif child.tag == qn('w:tbl'):
            items.append({'type': 'tbl', 'obj': doc.tables[ti]})
            ti += 1
    return items


def matches_heading(ptext, key):
    ptext = ptext.strip()
    if key in TOP_LEVEL_KEYS:
        return ptext.upper() == key
    if key == 'CAPITULO1':
        return ptext.upper().startswith('CAPÍTULO 1')
    m = NUM_HEADING_RE.match(ptext)
    return bool(m) and m.group(1) == key


# --------------------------------------------------------------------------
# Segment (paragraph/table run) rebuilding
# --------------------------------------------------------------------------

def process_editable_segment(existing_seq, md_blocks, para_template, insert_after_el, log):
    """existing_seq: list of body-index items (type 'p'/'tbl'). md_blocks: ('para',...)/('table',...).
    Reuses, creates or deletes paragraphs/tables so the region ends up matching md_blocks,
    while updating table cells in place. Returns the last xml element placed/kept."""
    ei = 0
    last_el = insert_after_el

    for block in md_blocks:
        kind = block[0]

        if kind == 'para':
            text = block[1]
            # skip forward past a mismatched table if one is next (shouldn't normally happen)
            if ei < len(existing_seq) and existing_seq[ei]['type'] == 'p':
                p = existing_seq[ei]['obj']
                rebuild_paragraph_text(p, text)
                last_el = p._p
                ei += 1
            else:
                new_p = clone_paragraph(para_template, text)
                last_el.addnext(new_p._p)
                last_el = new_p._p

        elif kind == 'table':
            rows = block[1]
            if ei < len(existing_seq) and existing_seq[ei]['type'] == 'tbl':
                table = existing_seq[ei]['obj']
                update_table(table, rows, log)
                last_el = table._tbl
                ei += 1
            else:
                log.append("    [WARN] markdown table has no matching docx table to update; skipped")

    # delete any leftover existing paragraphs/tables not consumed
    for j in range(ei, len(existing_seq)):
        item = existing_seq[j]
        el = item['obj']._p if item['type'] == 'p' else item['obj']._tbl
        el.getparent().remove(el)

    return last_el


def split_on_images(existing_seq):
    """Split existing_seq into alternating (editable_run, protected_group, editable_run, ...).
    Returns list of ('run', items) / ('protected', items)."""
    segments = []
    buf = []
    mode = None
    for item in existing_seq:
        is_img = item['type'] == 'p' and paragraph_has_image(item['obj'])
        this_mode = 'protected' if is_img else 'run'
        if mode is not None and this_mode != mode:
            segments.append((mode, buf))
            buf = []
        mode = this_mode
        buf.append(item)
    if buf:
        segments.append((mode, buf))
    return segments


def split_md_on_images(blocks):
    segments = []
    buf = []
    mode = None
    for b in blocks:
        this_mode = 'protected' if b[0] == 'image' else 'run'
        if mode is not None and this_mode != mode:
            segments.append((mode, buf))
            buf = []
        mode = this_mode
        buf.append(b)
    if buf:
        segments.append((mode, buf))
    return segments


def rebuild_region(existing_seq, blocks, para_template, insert_after_el, log):
    existing_segments = split_on_images(existing_seq)
    md_segments = split_md_on_images(blocks)

    existing_runs = [seg for mode, seg in existing_segments if mode == 'run']
    protected_groups = [seg for mode, seg in existing_segments if mode == 'protected']
    md_runs = [seg for mode, seg in md_segments if mode == 'run']

    last_el = insert_after_el

    if len(protected_groups) == 0:
        # simple case: no embedded images anywhere in this region
        flat_blocks = [b for seg in md_runs for b in seg] if md_runs else []
        last_el = process_editable_segment(existing_seq, flat_blocks, para_template, last_el, log)
        return last_el

    # region has one or more protected image groups; interleave editable runs
    # around them, leaving protected groups completely untouched in place.
    n_runs = max(len(existing_runs), len(md_runs), len(protected_groups) + 1)
    for i in range(n_runs):
        existing_run = existing_runs[i] if i < len(existing_runs) else []
        md_run = md_runs[i] if i < len(md_runs) else []
        last_el = process_editable_segment(existing_run, md_run, para_template, last_el, log)
        if i < len(protected_groups):
            grp = protected_groups[i]
            if grp:
                last_el = grp[-1]['obj']._p
                log.append(f"    [PROTECTED] left {len(grp)} image paragraph(s) untouched")

    return last_el


# --------------------------------------------------------------------------
# Main update routine
# --------------------------------------------------------------------------

def update_docx(original_docx, markdown_file, output_docx):
    log = []
    headings = parse_markdown(markdown_file)
    doc = Document(original_docx)

    body_index = build_body_index(doc)

    # locate anchor position (index into body_index) for each heading key, in order
    anchor_pos = {}
    cursor = 0
    for h in headings:
        found = None
        for pos in range(cursor, len(body_index)):
            item = body_index[pos]
            if item['type'] != 'p':
                continue
            if matches_heading(item['obj'].text, h.key):
                found = pos
                break
        anchor_pos[h.key] = found
        if found is not None:
            cursor = found + 1
        log.append(f"[ANCHOR] {h.key!r:14s} -> {'body_index ' + str(found) if found is not None else 'NOT FOUND (will insert)'}")

    # global fallback templates for brand-new paragraphs/headings
    numbered_heading_template = body_index[anchor_pos['1.1']]['obj']
    toplevel_heading_template = body_index[anchor_pos['RESUMEN']]['obj']
    # first body paragraph right after RESUMEN's heading is a plain, single-run
    # Normal-style paragraph -- safe generic template for any new body text
    body_para_template = body_index[anchor_pos['RESUMEN'] + 1]['obj']

    # process headings strictly in order, tracking an insertion cursor for
    # brand-new sections that have no existing anchor
    last_el = None
    for idx, h in enumerate(headings):
        pos = anchor_pos.get(h.key)
        next_pos = None
        for nh in headings[idx + 1:]:
            next_pos = anchor_pos.get(nh.key)
            if next_pos is not None:
                break

        if pos is not None:
            heading_para = body_index[pos]['obj']
            rebuild_paragraph_text(heading_para, h.raw_text)
            region_end = next_pos if next_pos is not None else len(body_index)
            existing_seq = body_index[pos + 1:region_end]
            template = numbered_heading_template  # unused here, kept for symmetry
            last_el = rebuild_region(existing_seq, h.blocks, body_para_template, heading_para._p, log)
            log.append(f"[SECTION] {h.key} updated in place ({len(existing_seq)} existing items -> {len(h.blocks)} md blocks)")
        else:
            # brand new section: create heading paragraph then its content
            if last_el is None:
                raise RuntimeError(f"Cannot insert heading {h.key!r}: no prior anchor to insert after")
            template = toplevel_heading_template if h.key in TOP_LEVEL_KEYS else numbered_heading_template
            new_heading = clone_paragraph(template, h.raw_text)
            last_el.addnext(new_heading._p)
            last_el = new_heading._p
            last_el = rebuild_region([], h.blocks, body_para_template, last_el, log)
            log.append(f"[SECTION] {h.key} created new (0 existing items -> {len(h.blocks)} md blocks)")

    doc.save(output_docx)
    return log


if __name__ == '__main__':
    original_docx = 'docs/reports/IDI - Primer Informe.docx'
    markdown_file = 'docs/reports/IDI_Capitulo1_v3.md'
    output_docx = 'docs/reports/IDI_Primer_Informe_Updated_Final.docx'

    try:
        log = update_docx(original_docx, markdown_file, output_docx)
        print('\n'.join(log))
        print(f"\n[OK] Saved to: {output_docx}")
    except Exception as e:
        print(f"\n[ERROR] {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
