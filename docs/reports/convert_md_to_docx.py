#!/usr/bin/env python
"""
Convert IDI_Capitulo1_v3.md to a properly formatted Word document
using python-docx with professional styling and formatting.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders in a table."""
    tcPr = cell._tcPr
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '12')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def parse_markdown(filepath):
    """Parse markdown file and return structured content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

def add_heading_with_style(doc, text, level):
    """Add heading with appropriate style."""
    if level == 1:
        p = doc.add_paragraph(text, style='Heading 1')
    elif level == 2:
        p = doc.add_paragraph(text, style='Heading 2')
    elif level == 3:
        p = doc.add_paragraph(text, style='Heading 3')
    else:
        p = doc.add_paragraph(text, style='Normal')
        p.runs[0].bold = True

    return p

def parse_table_markdown(table_text):
    """Parse markdown table into rows and columns."""
    lines = table_text.strip().split('\n')
    rows = []
    for i, line in enumerate(lines):
        if i == 1:  # Skip separator line
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows

def add_table_to_doc(doc, table_data):
    """Add a formatted table to document."""
    if not table_data:
        return

    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'

    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text.strip()

            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set header background color
                tcPr = cell._element.get_or_add_tcPr()
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '4472C4')
                tcPr.append(shading_elm)
            else:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def convert_markdown_to_docx(md_filepath, docx_filepath):
    """Convert markdown to Word document."""
    content = parse_markdown(md_filepath)
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split content by lines
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headers
        if line.startswith('# '):
            add_heading_with_style(doc, line[2:].strip(), 1)
            i += 1
        elif line.startswith('## '):
            add_heading_with_style(doc, line[3:].strip(), 2)
            i += 1
        elif line.startswith('### '):
            add_heading_with_style(doc, line[4:].strip(), 3)
            i += 1
        elif line.startswith('#### '):
            add_heading_with_style(doc, line[5:].strip(), 4)
            i += 1

        # Horizontal rules
        elif line.startswith('─') or line.startswith('----'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1

        # Tables
        elif line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            table_text = '\n'.join(table_lines)
            table_data = parse_table_markdown(table_text)
            add_table_to_doc(doc, table_data)
            doc.add_paragraph()  # Add space after table

        # Code blocks
        elif line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines), style='Normal')
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
            i += 1

        # Bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.lstrip('- * ').strip(), style='List Bullet')
            i += 1

        # Empty lines
        elif not line.strip():
            doc.add_paragraph()
            i += 1

        # Regular paragraphs
        else:
            p = doc.add_paragraph(line)
            i += 1

    # Save document
    doc.save(docx_filepath)
    print(f"[OK] Document saved to: {docx_filepath}")

if __name__ == '__main__':
    md_file = 'docs/reports/IDI_Capitulo1_v3.md'
    docx_file = 'docs/reports/IDI_Primer_Informe_v2.docx'
    convert_markdown_to_docx(md_file, docx_file)
